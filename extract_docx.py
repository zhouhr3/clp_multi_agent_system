#!/usr/bin/env python3
import docx
import sys

def extract_text_from_docx(file_path):
    """
    提取Word文档中的文本内容
    """
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        # 提取段落文本
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        return "\n\n".join(full_text)
    except Exception as e:
        return f"错误：无法读取文档 - {str(e)}"

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("用法: python extract_docx.py <docx文件路径>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    text_content = extract_text_from_docx(file_path)
    
    # 将内容输出到文件
    output_file = file_path.rsplit(".", 1)[0] + ".txt"
    with open(output_file, "w", encoding="utf-8") as f:
        f.write(text_content)
    
    print(f"文档内容已提取并保存到: {output_file}")
    print("\n内容预览:\n" + "-"*50)
    print(text_content[:1000] + "..." if len(text_content) > 1000 else text_content)
